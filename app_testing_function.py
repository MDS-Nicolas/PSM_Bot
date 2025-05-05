#############################################
# Extraction functions (asynchronous version)
#############################################
import os
import asyncio
from docx import Document
from dotenv import load_dotenv, find_dotenv

def load_api_key():
    """Load the API key from the .env file."""
    print("Api key loading...")
    load_dotenv(find_dotenv())
    api_key = os.getenv("API_KEY")
    if not api_key:
        raise ValueError("API_KEY is not set in your .env file.")
    print("API key loaded.")
    return api_key

async def process_pdf(openai_client, input_pdf_path, output_docx_path):
    """
    Dummy function that takes a PDF file and converts it to a DOCX file.
    
    Args:
        openai_client: OpenAI client instance
        input_pdf_path: Path to the input PDF file
        output_docx_path: Path to save the output DOCX file
    
    Returns:
        Document object after saving the DOCX file
    """
    print(f"Processing PDF file: {input_pdf_path}")
    print(f"Output will be saved to: {output_docx_path}")
    
    # In a real implementation, we would process the PDF content here
    # For this dummy function, we'll just create a simple DOCX file
    
    # Simulate processing time
    await asyncio.sleep(2)
    
    # Create a simple DOCX document
    doc = Document()
    doc.add_heading('PDF Conversion Result', 0)
    doc.add_paragraph(f'This is a dummy conversion of the PDF file: {os.path.basename(input_pdf_path)}')
    doc.add_paragraph('In a real implementation, this would contain the actual content extracted from the PDF.')
    
    # Save the document
    doc.save(output_docx_path)
    
    print(f"PDF processing complete. DOCX file saved to: {output_docx_path}")
    return doc

async def main_async():
    """Async main function to test the PDF processing."""
    try:
        # Load the API key
        api_key = load_api_key()
        
        # Initialize OpenAI client 
        from openai import AsyncOpenAI
        openai_client = AsyncOpenAI(api_key=api_key)
        
        # Test paths - replace with actual paths as needed
        input_pdf_path = "testingPDFshortTABLE.pdf"  # This doesn't need to exist for the dummy function
        output_docx_path = "app_output_result.docx"
        
        # Process the PDF
        doc = await process_pdf(openai_client, input_pdf_path, output_docx_path)
        
        # Example of using the returned Document object
        print("Adding additional content to the document...")
        doc.add_paragraph("This paragraph was added after the initial processing.")
        doc.save("modified_output.docx")
        print("Modified document saved to: modified_output.docx")
        
    except Exception as e:
        print(f"Error in main: {e}")

def main():
    """Run the async main function."""
    asyncio.run(main_async())

if __name__ == "__main__":
    main()